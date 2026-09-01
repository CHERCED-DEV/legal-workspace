# -*- coding: utf-8 -*-
"""Conversor Markdown -> Word con tablas de verdad.

Puerto fiel de md2docx.js, que dependia de Node y de una ruta NODE_PATH fija
a la maquina de quien lo escribio. ADR-018: la oficina es Python, un solo
tiempo de ejecucion. ADR-014 decision 2: toda tabla prometida "lista para
pegar" es una tabla de Word real.

    python md2docx.py entrada.md salida.docx ["Titulo"] ["Subtitulo"]

Si se fuerza el subtitulo, el original NO se pierde: baja al cuerpo como
bloque destacado. En la primera version del conversor desaparecia sin dejar
rastro, y esa linea suele ser el descargo de la salida.
"""
import io
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, Twips, RGBColor
except ImportError:
    sys.stderr.write(
        "FALTA python-docx. El comando puede seguir sin el: entrega el contenido\n"
        "en texto y declara que no se pudo producir el .docx (ADR-014 decision 4).\n"
        "  pip install python-docx\n")
    raise SystemExit(3)

# --- medidas, identicas a las del conversor anterior -----------------------
ANCHO, ALTO = 12240, 15840          # carta, en dxa
MARGEN = 1080
CONTENIDO = ANCHO - 2 * MARGEN      # 10080
FUENTE = "Calibri"
TINTA = "1F3864"                    # azul de encabezados y cabecera de tabla
GRIS = "595959"
CREMA = "FFF2CC"                    # cajas
CEBRA = "F2F2F2"


# --- utilidades de XML que python-docx no expone --------------------------
def _sombrear(celda, relleno):
    pr = celda._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')        # nunca 'solid': tapa el texto en Word
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), relleno)
    pr.append(shd)


def _margenes_celda(tabla, top=80, bottom=80, left=110, right=110):
    pr = tabla._tbl.tblPr
    mar = OxmlElement('w:tblCellMar')
    for etiqueta, valor in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        e = OxmlElement('w:' + etiqueta)
        e.set(qn('w:w'), str(valor))
        e.set(qn('w:type'), 'dxa')
        mar.append(e)
    pr.append(mar)


def _borde_inferior(parrafo, grosor, color, espacio):
    pr = parrafo._p.get_or_add_pPr()
    bdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), str(grosor))       # octavos de punto
    b.set(qn('w:space'), str(espacio))
    b.set(qn('w:color'), color)
    bdr.append(b)
    pr.append(bdr)


def _repetir_cabecera(fila):
    pr = fila._tr.get_or_add_trPr()
    e = OxmlElement('w:tblHeader')
    e.set(qn('w:val'), 'true')
    pr.append(e)


# --- texto con **negrita** y *cursiva* ------------------------------------
_MARCAS = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*)')


def _runs(parrafo, texto, medio=21, color=None):
    """medio: tamano en medios-punto, como lo expresaba el conversor anterior."""
    def anadir(t, negrita=False, cursiva=False):
        if not t:
            return
        r = parrafo.add_run(t)
        r.font.name = FUENTE
        r.font.size = Pt(medio / 2.0)
        r.bold = negrita
        r.italic = cursiva
        if color:
            r.font.color.rgb = RGBColor.from_string(color)

    ultimo = 0
    for m in _MARCAS.finditer(texto):
        anadir(texto[ultimo:m.start()])
        t = m.group(0)
        if t.startswith('**'):
            anadir(t[2:-2], negrita=True)
        else:
            anadir(t[1:-1], cursiva=True)
        ultimo = m.end()
    anadir(texto[ultimo:])


# --- bloques ---------------------------------------------------------------
def p(doc, texto, despues=140, sangria=0, alineacion=WD_ALIGN_PARAGRAPH.JUSTIFY,
      medio=21, color=None, contenedor=None):
    par = (contenedor or doc).add_paragraph()
    par.alignment = alineacion
    f = par.paragraph_format
    f.space_after = Twips(despues)
    f.line_spacing = Pt(280 / 20.0)
    if sangria:
        f.left_indent = Twips(sangria)
    _runs(par, texto, medio, color)
    return par


def h(doc, texto, nivel):
    par = doc.add_paragraph(style='Heading %d' % nivel)
    f = par.paragraph_format
    f.space_before = Twips(260)
    f.space_after = Twips(130)
    medio = {1: 30, 2: 25, 3: 22}[nivel]
    r = par.add_run(texto)
    r.font.name = FUENTE
    r.font.size = Pt(medio / 2.0)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(TINTA)
    return par


def vineta(doc, texto):
    par = doc.add_paragraph(style='List Bullet')
    f = par.paragraph_format
    f.space_after = Twips(70)
    f.line_spacing = Pt(14)
    _runs(par, texto, 21)
    return par


def caja(doc, lineas, relleno=CREMA):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    _margenes_celda(t, 120, 120, 150, 150)
    c = t.cell(0, 0)
    c.width = Twips(CONTENIDO)
    t.columns[0].width = Twips(CONTENIDO)
    _sombrear(c, relleno)
    c.paragraphs[0]._p.getparent().remove(c.paragraphs[0]._p)
    for l in lineas:
        p(doc, l, despues=60, contenedor=c)
    doc.add_paragraph()
    return t


def tabla(doc, cabeceras, filas, pesos):
    total = sum(pesos) or 1
    cols = [int(round(CONTENIDO * w / total)) for w in pesos]
    cols[-1] += CONTENIDO - sum(cols)          # el redondeo se paga en la ultima

    t = doc.add_table(rows=1 + len(filas), cols=len(cabeceras))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    _margenes_celda(t)

    # Doble anchura: en la rejilla y en cada celda. Con una sola, Word reparte
    # a su gusto y las tablas salen descuadradas.
    for i, w in enumerate(cols):
        t.columns[i].width = Twips(w)

    for i, texto in enumerate(cabeceras):
        c = t.cell(0, i)
        c.width = Twips(cols[i])
        _sombrear(c, TINTA)
        par = c.paragraphs[0]
        par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        par.paragraph_format.space_after = Twips(0)
        r = par.add_run(str(texto))
        r.font.name = FUENTE
        r.font.size = Pt(9.5)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string("FFFFFF")
    _repetir_cabecera(t.rows[0])

    for j, fila in enumerate(filas):
        for i, texto in enumerate(fila):
            c = t.cell(j + 1, i)
            c.width = Twips(cols[i])
            _sombrear(c, CEBRA if j % 2 else "FFFFFF")
            par = c.paragraphs[0]
            par.alignment = WD_ALIGN_PARAGRAPH.LEFT
            f = par.paragraph_format
            f.space_after = Twips(0)
            f.line_spacing = Pt(12.5)
            _runs(par, str(texto), 19)
    doc.add_paragraph()
    return t


def documento(titulo, subtitulo):
    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Twips(ANCHO), Twips(ALTO)
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Twips(MARGEN)
    n = doc.styles['Normal']
    n.font.name = FUENTE
    n.font.size = Pt(10.5)

    par = doc.add_paragraph()
    par.paragraph_format.space_after = Twips(60)
    r = par.add_run(titulo)
    r.font.name = FUENTE
    r.font.size = Pt(17)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(TINTA)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Twips(200)
    r = sub.add_run(subtitulo)
    r.font.name = FUENTE
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(GRIS)
    _borde_inferior(sub, 12, TINTA, 8)
    return doc


# --- analisis del Markdown -------------------------------------------------
def limpia(s):
    s = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', s)   # enlaces -> solo texto
    s = re.sub(r'`([^`]+)`', r'\1', s)               # codigo en linea
    return s.rstrip()


def anchos(cabeceras, filas):
    n = len(cabeceras)
    w = [0] * n
    for r in [cabeceras] + filas:
        for i in range(n):
            w[i] = max(w[i], min(len(r[i] or ''), 90))
    return [max(x, 6) for x in w]


RE_FICHA = re.compile(r'^\*\*[HFVC]-\d+[^*]*\*\*\s*$')
RE_CAMPO_PLANO = re.compile(r'^\s*[-*·]\s+[A-ZÁÉÍÓÚÑa-z][^:]{1,24}:\s')
RE_CAMPO_NEGRITA = re.compile(r'^\s*[-*·]\s+\*\*[^*]+\*\*')
RE_VINETA = re.compile(r'^\s*[-*·]\s+')
RE_NUMERADA = re.compile(r'^\s*\d+\.\s+')


def _continuacion(lineas, j, sangria=2):
    """Une las lineas indentadas que continuan un elemento de lista."""
    extra = []
    while (j + 1 < len(lineas) and re.match(r'^\s{%d,}\S' % sangria, lineas[j + 1])
           and not re.match(r'^\s*[-*·|>#]', lineas[j + 1])):
        j += 1
        extra.append(limpia(lineas[j]).strip())
    return j, extra


def convierte(doc, md):
    lineas = md.split('\n')
    lineas = [l.rstrip('\r') for l in lineas]
    i, bloques = 0, 0
    titulo, subtitulo = None, None
    pendiente_subtitulo = []

    while i < len(lineas):
        L = limpia(lineas[i])

        if titulo is None and re.match(r'^#\s+', L):
            titulo = re.sub(r'^#\s+', '', L)
            i += 1
            while i < len(lineas) and not lineas[i].strip():
                i += 1
            if i < len(lineas) and lineas[i].strip() and not re.match(r'^[#>|-]', lineas[i]):
                subtitulo = limpia(lineas[i]).replace('**', '')
                i += 1
            continue

        if not L.strip() or re.match(r'^---+$', L.strip()):
            i += 1
            continue

        for marca, nivel in (('####', 3), ('###', 2), ('##', 1), ('#', 1)):
            if re.match(r'^%s\s+' % re.escape(marca), L):
                h(doc, re.sub(r'^%s\s+' % re.escape(marca), '', L), nivel)
                bloques += 1
                break
        else:
            if re.match(r'^>\s?', L):                                  # cita -> caja
                buf = []
                while i < len(lineas) and re.match(r'^>\s?', lineas[i]):
                    t = re.sub(r'^>\s?', '', limpia(lineas[i]))
                    if t.strip():
                        buf.append(t)
                    i += 1
                caja(doc, buf or [' '])
                bloques += 1
                continue

            if re.match(r'^\s*\|', L):                                 # tabla
                crudas = []
                while i < len(lineas) and re.match(r'^\s*\|', lineas[i]):
                    crudas.append(limpia(lineas[i]).strip())
                    i += 1
                celdas = [[c.strip() for c in f.strip('|').split('|')] for f in crudas]
                cuerpo = [r for r in celdas
                          if not all(re.match(r'^:?-{2,}:?$', c) or c == '' for c in r)]
                if cuerpo:
                    hd, rs = cuerpo[0], cuerpo[1:]
                    n = len(hd)
                    norm = [(r[:n] + [''] * (n - len(r[:n]))) for r in rs]
                    tabla(doc, hd, norm, anchos(hd, norm))
                    bloques += 1
                continue

            if RE_FICHA.match(L.strip()):                              # ficha -> encabezado
                h(doc, L.strip().strip('*'), 2)
                bloques += 1
                i += 1
                continue

            # lista "- Campo: valor" -> tabla de dos columnas
            if RE_CAMPO_PLANO.match(L) and '**' not in L.split(':')[0]:
                filas, j = [], i
                while (j < len(lineas) and RE_CAMPO_PLANO.match(lineas[j])
                       and '**' not in lineas[j].split(':')[0]):
                    t = re.sub(r'^\s*[-*·]\s+', '', limpia(lineas[j]))
                    j, extra = _continuacion(lineas, j)
                    t = ' '.join([t] + extra)
                    k = t.index(':')
                    filas.append([t[:k], t[k + 1:].strip()])
                    j += 1
                if len(filas) >= 3:
                    tabla(doc, ['Campo', 'Contenido'], filas, [18, 82])
                    bloques += 1
                    i = j
                    continue

            # lista "- **Campo:** valor" -> tabla de dos columnas
            if RE_CAMPO_NEGRITA.match(L):
                filas, j = [], i
                while j < len(lineas) and RE_CAMPO_NEGRITA.match(lineas[j]):
                    t = re.sub(r'^\s*[-*·]\s+', '', limpia(lineas[j]))
                    j, extra = _continuacion(lineas, j)
                    t = ' '.join([t] + extra)
                    m = re.match(r'^\*\*([^*]+?)\*\*[:：]?\s*(.*)$', t)
                    filas.append([re.sub(r'[:：]$', '', m.group(1)), m.group(2)] if m else ['', t])
                    j += 1
                if len(filas) >= 3:
                    tabla(doc, ['Campo', 'Contenido'], filas, [22, 78])
                    bloques += 1
                    i = j
                    continue

            if RE_VINETA.match(L):                                     # vinetas
                while i < len(lineas) and RE_VINETA.match(lineas[i]):
                    t = re.sub(r'^\s*[-*·]\s+', '', limpia(lineas[i]))
                    i, extra = _continuacion(lineas, i)
                    vineta(doc, ' '.join([t] + extra))
                    bloques += 1
                    i += 1
                continue

            if RE_NUMERADA.match(L):                                   # numeradas
                while i < len(lineas) and RE_NUMERADA.match(lineas[i]):
                    t = limpia(lineas[i]).strip()
                    extra = []
                    while (i + 1 < len(lineas) and re.match(r'^\s{3,}\S', lineas[i + 1])
                           and not re.match(r'^\s*[-*·|>#]', lineas[i + 1])
                           and not RE_NUMERADA.match(lineas[i + 1])):
                        i += 1
                        extra.append(limpia(lineas[i]).strip())
                    p(doc, ' '.join([t] + extra), sangria=280)
                    bloques += 1
                    i += 1
                continue

            buf = [L]                                                  # parrafo
            i += 1
            while (i < len(lineas) and lineas[i].strip()
                   and not re.match(r'^\s*[|>#]', lineas[i])
                   and not re.match(r'^\s*[-*·]\s', lineas[i])
                   and not RE_NUMERADA.match(lineas[i])
                   and not re.match(r'^---+$', lineas[i].strip())):
                buf.append(limpia(lineas[i]))
                i += 1
            p(doc, ' '.join(buf))
            bloques += 1
            continue
        i += 1

    return titulo or 'Documento', subtitulo or '', bloques


def main():
    if len(sys.argv) < 3:
        sys.stderr.write(__doc__ + '\n')
        raise SystemExit(2)
    entrada, salida = Path(sys.argv[1]), Path(sys.argv[2])
    titulo_forzado = sys.argv[3] if len(sys.argv) > 3 else None
    sub_forzado = sys.argv[4] if len(sys.argv) > 4 else None

    md = io.open(entrada, encoding='utf-8').read()

    # Dos pasadas: la primera solo para saber titulo y subtitulo, porque el
    # encabezado va antes que el cuerpo y el cuerpo puede rescatar el subtitulo.
    sonda = Document()
    tit, sub, _ = convierte(sonda, md)

    doc = documento(titulo_forzado or tit, sub_forzado or sub)
    # Si se fuerza el subtitulo, el original NO se pierde: baja como descargo.
    if sub_forzado and sub.strip():
        caja(doc, [sub])
    _, _, bloques = convierte(doc, md)
    doc.save(salida)

    tablas = len(doc.tables)
    filas = sum(len(t.rows) for t in doc.tables)
    print('OK  %s  -  %d bloques, %d tablas reales, %d filas'
          % (salida.name, bloques, tablas, filas))


if __name__ == '__main__':
    main()
