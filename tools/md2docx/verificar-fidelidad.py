#!/usr/bin/env python
"""
Control de fidelidad Markdown -> Word.

Comprueba que la conversion no perdio contenido. NO es una comparacion exacta:
mide que fraccion de las palabras del .md aparece en el .docx, ignorando la
puntuacion de Markdown.

Por que existe: en el primer uso real, forzar el titulo desde la linea de
comandos hacia desaparecer la linea de descargo que iba debajo del titulo
("Propuesta para su revision. Nada de esto esta comprobado por ningun
sistema."). Se veia bien y faltaba justo la frase que no puede faltar. Este
script lo detecto; una lectura por encima no lo habria detectado.

Uso:
    python verificar-fidelidad.py salida.docx entrada.md
    python verificar-fidelidad.py --pares lista.txt      (una linea "docx<TAB>md")

Umbrales:
    >= 99 %   ok
    95-99 %   REVISAR  (suele ser ruido de la propia medicion: etiquetas con
                        dos puntos que en el Word quedan sin ellos)
    < 95 %    PERDIDA  (hay que mirar que falta, de verdad)

Requiere: python-docx
"""
import re
import sys


def normaliza(texto: str) -> str:
    """Quita la puntuacion propia de Markdown y colapsa espacios."""
    texto = re.sub(r"[*`|>#_-]", " ", texto)
    return re.sub(r"\s+", " ", texto).strip().lower()


def texto_del_docx(ruta: str) -> str:
    from docx import Document

    doc = Document(ruta)
    partes = [p.text for p in doc.paragraphs]
    partes += [c.text for t in doc.tables for f in t.rows for c in f.cells]
    return normaliza(" ".join(partes))


def compara(ruta_docx: str, ruta_md: str):
    origen = normaliza(open(ruta_md, encoding="utf-8").read())
    destino = texto_del_docx(ruta_docx)

    # Solo palabras largas: las cortas generan ruido y no distinguen nada.
    palabras = [p for p in set(origen.split()) if len(p) > 5]
    if not palabras:
        return 100.0, []
    faltan = sorted(p for p in palabras if p not in destino)
    retencion = 100.0 * (len(palabras) - len(faltan)) / len(palabras)
    return retencion, faltan


def veredicto(retencion: float) -> str:
    if retencion >= 99:
        return "ok"
    if retencion >= 95:
        return "REVISAR"
    return "PERDIDA"


def main(argv):
    if len(argv) == 3 and argv[1] == "--pares":
        pares = [
            linea.rstrip("\n").split("\t")
            for linea in open(argv[2], encoding="utf-8")
            if linea.strip()
        ]
    elif len(argv) == 3:
        pares = [(argv[1], argv[2])]
    else:
        print(__doc__)
        return 2

    problemas = 0
    for ruta_docx, ruta_md in pares:
        retencion, faltan = compara(ruta_docx, ruta_md)
        marca = veredicto(retencion)
        nombre = ruta_docx.split("\\")[-1].split("/")[-1]
        print(f"{nombre[:52]:54s} {retencion:5.1f}%  {marca}")
        if marca != "ok":
            print("    faltan:", ", ".join(faltan[:25]))
            if marca == "PERDIDA":
                problemas += 1
    return 1 if problemas else 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))
